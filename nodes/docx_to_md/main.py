"""
DOCX → 마크다운 변환 노드.

구현 우선순위:
1. kordoc CLI (설치되어 있으면)
2. python-docx (fallback)
"""

import shutil
import subprocess
from pathlib import Path


def _convert_with_kordoc(docx_path: str) -> str | None:
    """kordoc CLI로 변환 시도."""
    # Windows에서 CreateProcess가 npx.cmd를 실행하려면 확장자 포함 전체 경로 필요
    npx = shutil.which("npx")
    if not npx:
        return None

    try:
        cmd = [npx, "kordoc", docx_path, "--format", "json"]
        result = subprocess.run(
            cmd, capture_output=True, text=True,
            encoding="utf-8", errors="replace", timeout=120,
        )
        if result.returncode == 0 and result.stdout.strip():
            import json
            parsed = json.loads(result.stdout)
            md = parsed.get("markdown") if isinstance(parsed, dict) else None
            if isinstance(md, str) and md.strip():
                return md
    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        pass
    return None


def _convert_with_fallback(docx_path: str) -> str:
    """python-docx로 DOCX 텍스트 추출."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx 미설치. pip install python-docx 후 다시 시도하세요."
        )

    doc = Document(docx_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 스타일 기반 마크다운 변환
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            parts.append(f"# {text}")
        elif "heading 2" in style_name:
            parts.append(f"## {text}")
        elif "heading 3" in style_name:
            parts.append(f"### {text}")
        elif "heading 4" in style_name:
            parts.append(f"#### {text}")
        elif "list" in style_name:
            parts.append(f"- {text}")
        else:
            parts.append(text)

    # 표 추출
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if rows:
            # 헤더 구분선 추가
            header = rows[0]
            col_count = header.count("|") - 1
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            table_md = rows[0] + "\n" + separator
            if len(rows) > 1:
                table_md += "\n" + "\n".join(rows[1:])
            parts.append("\n" + table_md + "\n")

    if parts:
        return "\n\n".join(parts)
    raise RuntimeError("DOCX에서 텍스트를 추출할 수 없음")


def execute(inputs: dict, params: dict, context: dict) -> dict:
    docx_path = inputs["파일"]

    if not Path(docx_path).exists():
        raise FileNotFoundError(f"DOCX 파일 없음: {docx_path}")

    context["progress"](0.1)
    context["log"]("DOCX 변환 시작")

    # kordoc 먼저 시도, 실패하면 python-docx
    result = _convert_with_kordoc(docx_path)
    if result is None:
        context["log"]("kordoc 미설치, python-docx로 변환")
        result = _convert_with_fallback(docx_path)

    context["progress"](1.0)
    context["log"](f"변환 완료 ({len(result)} 글자)")

    return {"텍스트": result}
