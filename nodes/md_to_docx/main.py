"""
마크다운 → DOCX 변환 노드.

python-docx로 마크다운 텍스트를 DOCX(Word) 파일로 변환한다.
헤딩, 단락, 표, 리스트를 지원한다.
"""

import os
import re


def _strip_inline(text: str) -> str:
    """볼드/이탤릭 마커 제거 (간이). 헤딩·리스트·표·인용·단락에 일관 적용."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def _add_table_to_doc(doc, table_lines: list[str]):
    """마크다운 표를 DOCX 표로 추가."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # 구분선(---|---) 건너뛰기
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")

    for i, row_data in enumerate(rows):
        for j in range(num_cols):
            text = row_data[j] if j < len(row_data) else ""
            table.rows[i].cells[j].text = _strip_inline(text)


def execute(inputs: dict, params: dict, context: dict) -> dict:
    md_text = inputs["텍스트"]
    output_name = params.get("output_name", "output")
    output_path = os.path.join(context["temp_dir"], f"{output_name}.docx")

    from docx import Document
    from docx.shared import Pt

    context["progress"](0.1)
    context["log"]("DOCX 변환 시작")

    doc = Document()

    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 헤딩
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = _strip_inline(heading_match.group(2).strip())
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # 구분선
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph("").paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # HTML 주석
        if line.strip().startswith("<!--"):
            i += 1
            continue

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # 표
        if re.match(r"^\|.*\|$", line.strip()):
            table_lines = []
            while i < len(lines) and re.match(r"^\|.*\|$", lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            _add_table_to_doc(doc, table_lines)
            continue

        # 리스트 항목
        list_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if list_match:
            text = _strip_inline(list_match.group(2).strip())
            p = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # 번호 리스트
        num_list_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if num_list_match:
            text = _strip_inline(num_list_match.group(2).strip())
            p = doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # 인용
        if line.strip().startswith("> "):
            text = _strip_inline(line.strip()[2:])
            p = doc.add_paragraph(text)
            p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else None
            i += 1
            continue

        # 코드 블록
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 닫는 ``` 건너뛰기
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # 일반 단락
        clean_text = _strip_inline(line.strip())
        doc.add_paragraph(clean_text)
        i += 1

    doc.save(output_path)

    context["progress"](1.0)
    context["log"](f"DOCX 생성 완료: {output_path}")

    return {"파일": output_path}
